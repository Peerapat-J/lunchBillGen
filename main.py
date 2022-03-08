from ast import Global
from click import style

from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import Length
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

from  gui import *
import os
import string

#curDir = os.getcwd()
#msWord = Document()
#msWord.save("launchBill01.docx")

document = Document()

#
#   define page layout
#
sections = document.sections
for section in sections:
    section.top_margin = Cm(0.4)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    
#doc = docx.Document('test.docx')
#print(document.paragraphs[0].text)

#
#   This need to fix as in front of text
#
paragraph = document.add_picture("ตราครุฑ.jpeg", width=Cm(1.6), height=Cm(1.7))

#
#   It doesn't work with pic ?
#
#paragraph.alignment = 1

paragraph = document.add_paragraph()
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)

#
#   Note: alignment object mode.
#       0: left
#       1: center
#       2: right
#       3: justify
#       etc: https://python-docx.readthedocs.io/en/latest/api/enum/WdAlignParagraph.html
#
paragraph_format.alignment = 1 # center

paragraph = paragraph.add_run('บันทึกข้อความ')
paragraph.font.bold = True
paragraph.font.name = 'TH SarabunPSK'
paragraph.font.size = Pt(16)

#paragraph.font.bold = True
#paragraph.add_run('_____END')
#paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

#
#   Global font define.
#
style = document.styles['Normal']
#style.paragraph_format.line_spaceing = 0
font = style.font
font.name = 'TH SarabunPSK'
font.size = Pt(16)
#font.bold = True

#
# Section define.
#
paragraph = document.add_paragraph()
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = 0 # left

paragraph = paragraph.add_run('ส่วนราชการ โรงเรียนวัดสระแก้ว')
paragraph.font.bold = False
paragraph.font.name = 'TH SarabunPSK'
paragraph.font.size = Pt(16)

#
#   document number.
#
paragraph = document.add_paragraph()
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = 0 # left

#
#   note: need to be dynamic not static hardcode like this, will fix this latter.
#
paragraph = paragraph.add_run('ที่  ............. /2564\t\tวันที่ 28 กุมภาพันธ์ 2565')
paragraph.font.bold = False
paragraph.font.name = 'TH SarabunPSK'
paragraph.font.size = Pt(16)

#
#   subject.
#
paragraph = document.add_paragraph()
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = 0 # left

paragraph = paragraph.add_run('เรื่อง\tขออนุมัติการจ่ายเงิน')
paragraph.font.bold = False
paragraph.font.name = 'TH SarabunPSK'
paragraph.font.size = Pt(16)

#   line object.
#paragraph = document.add_paragraph('')
#paragraph_format = paragraph.paragraph_format
#paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
#paragraph_format.space_before = Pt(0)
#paragraph_format.space_after = Pt(0)
#paragraph_format.alignment = 0
#paragraph = paragraph.add_run('____________________________________________________________________________')
#paragraph.font.bold = False
#paragraph.font.name = 'TH SarabunPSK'
#paragraph.font.size = Pt(16)

#
#   This need to fix as in front of text
#
#paragraph = document.add_picture("line.png", width=Cm(17))

#
#   start subject detail.
#
paragraph = document.add_paragraph('______________________________________________________________________________________')
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = 0 # left

#
#   start subject detail.
#
paragraph = document.add_paragraph()
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = 0 # left

paragraph = paragraph.add_run('เรียน  ผู้อำนวยการโรงเรียนวัดสระแก้ว')
paragraph.font.bold = False

#
#   start main detail section 1.
#
paragraph = document.add_paragraph('')
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = 0 # left

paragraph.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
paragraph.add_run('\tตามประกาศกระทรวงศึกษาธิการ เรื่อง หลักเกณฑ์การเปิดโรงเรียนหรือสถาบันการศึกษาตามข้อกำหนด ')
paragraph.add_run('ออกตามความในมาตรา ๙ แห่งพระราชกำหนด การบริหารราชการในสถานการณ์ฉุกเฉิน พ.ศ. ๒๕๔๘ ')
paragraph.add_run('(ฉบับที่ ๓๔) และมติคณะรัฐมนตรีเมื่อวันที่ 19 พฤษภาคม 2564 เรื่อง แนวทางการดำเนินโครงการอาหารเสริม ')
paragraph.add_run('(นม) โรงเรียนและการสนับสนุนอาหารกลางวันในโรงเรียนรองรับสถานการณ์การแพร่ระบาดของโรคติดเชื้อไวรัส ')
paragraph.add_run('โคโรนา 2019 (โรคโควิด 19) นั้น เพื่อให้การบริหารจัดการ เวลาเรียนและโครงการอาหารกลางวันสอดคล้องกับ ')
paragraph.add_run('ประกาศกระทรวงศึกษาธิการและมติคณะรัฐมนตรีดังกล่าว โรงเรียนวัดสระแก้วจึงจัดการเรียนการสอนออนไลน์ใน ')

#
#   need to be selected option in this part.
#
paragraph.add_run('วันที่ 1 กุมภาพันธ์ 2565 ถึง 28 กุมภาพันธ์ 2565 ')
paragraph.add_run('รวม 19 วัน ')
paragraph.add_run('และได้จัดสรรเงินอุดหนุนเป็นค่าอาหารกลางวัน ให้นักเรียนตั้งแต่ระดับชั้นอนุบาลปีที่ 2 ถึงระดับชั้นประถมศึกษาปีที่ 6 ')
paragraph.add_run('จำนวน 359 คน ')
paragraph.add_run('คนละ 399 บาท ')

#
#   start main detail section 2.
#
paragraph = document.add_paragraph('')
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = 0 # left
paragraph.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

#
#   need to have a space at the end to prevent docx content messup for some reason.
#
paragraph.add_run('\tจึงขออนุมัติเบิกเงินเพื่อจ่ายให้นักเรียนในวันที่ ')
#   need fix
paragraph.add_run('28 เดือน กุมภาพันธ์ พ.ศ. 2565 ')
paragraph.add_run('เป็นเงิน ')
paragraph.add_run('143,241 บาท ')
#   need converting number to thai word.
paragraph.add_run('(หนึ่งแสนสี่หมื่นสามพันสองร้อยสี่สิบเอ็ดบาทถ้วน)  ')

#
#   start main detail section 3.
#
paragraph = document.add_paragraph('\tจึงเรียนมาเพื่อโปรดพิจารณา')
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = 0 # left

#
#   start main detail section 3.1
#
paragraph = document.add_paragraph('\t1. อนุมัติให้จ่ายเงิน')
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = 0 # left

#
#   start main detail section 3.2
#
paragraph = document.add_paragraph('\t2. ')
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = 0 # left
paragraph.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
paragraph.add_run('แต่งตั้งคณะกรรมการจ่ายเงินตามหลักเกณฑ์การใช้จ่ายงบประมาณงบเงินอุดหนุนของสำนักงาน ')
paragraph.add_run('คณะกรรมการการศึกษาขั้นพื้นฐาน ดังนี้ ')


#
#   start main detail section 3.2.1
#
paragraph = document.add_paragraph('\t\t2.1 นายสำเนา สำราญจิตร์\tตำแหน่ง ครูชำนาญการพิเศษ ')
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = 0 # left

#
#   start main detail section 3.2.2
#
paragraph = document.add_paragraph('\t\t2.2 นางเกศราภรณ์ เกตุชาติ\tตำแหน่ง ครูชำนาญการ ')
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = 0 # left

#
#   start main detail section 3.2.3
#
paragraph = document.add_paragraph('\t\t2.3 นางจารุเนตร ทองจันดี\tตำแหน่ง ครูชำนาญการพิเศษ ')
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = 0 # left

#
#   Empty 1, refer from original document form
#
paragraph = document.add_paragraph()
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)

#
#   Head financial account signature part
#
paragraph = document.add_paragraph('\t\t\t\t\tลงชื่อ.................................................เจ้าหน้าที่การเงิน ')
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = 0  # left

#
#   Head financial account full name
#
paragraph = document.add_paragraph('\t\t\t\t\t      ( นางสาวจุฑามาศ ตั้งคงเจริญชัย ) ')
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = 0  # left

#
#   Head financial account sign date
#
paragraph = document.add_paragraph('\t\t\t\t\t\t28 กุมภาพันธ์ 2565 ')
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = 0  # left

#
#   Empty 2, refer from original document form
#
paragraph = document.add_paragraph()
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)

#
#   Comment section from above title
#
paragraph = document.add_paragraph('\t\t\t\t\tความเห็นของผู้บังคับบัญชา')
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)

#
#   Comment 1
#
paragraph = document.add_paragraph('\t\t\t\t\t\t1. ทราบ')
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)

#
#   Comment 2
#
paragraph = document.add_paragraph('\t\t\t\t\t\t2. อนุมัติ')
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)

#
#   Empty 3, refer from original document form
#
'''
paragraph = document.add_paragraph()
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
'''

#
#   Director signature
#
paragraph = document.add_paragraph('\t\t\t\t\tลงชื่อ.................................................ผู้อำนวยการโรงเรียนวัดสระแก้ว ')
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = 0  # left

#
#   Director full name
#
paragraph = document.add_paragraph('\t\t\t\t\t\t( นางภัสนันท์ รักษ์เมือง ) ')
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = 0  # left

#
#   Director sign date
#
paragraph = document.add_paragraph('\t\t\t\t\t\t   28 กุมภาพันธ์ 2565 ')
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = 0  # left

#
#   example from stackoverflow
#
#para = document.add_paragraph().add_run('GeeksforGeeks is a Computer Science portal for geeks.')
# Increasing size of the font
#para.font.name = 'TH SarabunPSK'
#para.font.size = Pt(12)
#para.font.bold = True

#
#   Done
#   Need to add some file name format like _d/m/y
#
fileName = 'บันทึกข้อความอาหารกลางวันประถม(new).docx'
document.save(fileName)

#
#   Convert to pdf ?
#
convert(fileName)

#
#   note to me
#
"""endCheck = True
i = 0
while(endCheck):
    print('i = ' + str(i))
    docCheck = document.paragraphs[i].text
    print("paragraphs[%2d] = "%(i), docCheck)
    if (docCheck == 'end'):
        endCheck = False
    else:
        i += 1
"""